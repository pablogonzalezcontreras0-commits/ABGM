import sys
from docx import Document
from docx.oxml import OxmlElement
import copy

def replace_text_in_paragraphs(paragraphs, replacements):
    for p in paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                # Simple run replacement
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                # If still there (split across runs)
                if key in p.text:
                    full_text = p.text.replace(key, value)
                    p.clear()
                    p.add_run(full_text)

def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, replacements)

def insert_table_after(paragraph, table):
    p = paragraph._p
    p.addnext(table._tbl)

def create_req_table(doc, template_table, req):
    # Copy table
    new_tbl = copy.deepcopy(template_table._tbl)
    # We need to wrap it in a Table object to modify it
    from docx.table import Table
    new_table = Table(new_tbl, doc)
    
    # Fill data
    # Row 0: ['Código', '{ Código }', 'Nombre', '{ Nombre }']
    new_table.rows[0].cells[1].text = req['codigo']
    new_table.rows[0].cells[3].text = req['nombre']
    # Row 1: ['Descripción', '{ Descripción...}', '{ Descripción...}', '{ Descripción...}']
    # It seems cells 1,2,3 are merged in the template. We just set text to cell 1
    new_table.rows[1].cells[1].text = req['desc']
    # Row 2: ['Prioridad', '{ 1 a 5 }', 'Rel.', '{ Código de Req/s. relacionado/s }']
    new_table.rows[2].cells[1].text = str(req['prioridad'])
    new_table.rows[2].cells[3].text = req['rel']
    # Row 3: ['Trazabilidad', '{ OPCIONAL... }', 'Ref.', '{ Doc/s. Relacionados/Aclaratorios }']
    new_table.rows[3].cells[1].text = req['trazabilidad']
    new_table.rows[3].cells[3].text = req['ref']
    
    return new_tbl

def main(input_path, output_path):
    doc = Document(input_path)

    # 1. Global replacements
    replacements = {
        "{ NOMBRE DEL PROYECTO }": "Sistema de Monitoreo para Reconectador Automático OSM27",
        "{ Nombre del proyecto }": "Sistema de monitoreo para reconectador automático OSM27",
        "{ NOMBRE EMPRESA }": "IPSEP",
        "{ BENEFICIARIO }": "UNRC",
        "{ uso y/o distribución }": "uso interno y académico",
        "{ convocatoria / protocolo de trabajo / convenio / etc }": "Trabajo Final de Aplicaciones TCP/IP",
        "{ institución / organización / empresa BENEFICIARIA }": "UNRC",
        "{ explicar las tareas }": "adquisición, almacenamiento y visualización de telemetría Modbus TCP",
        "{ Hardware / Firmware / Software }": "Software y Hardware",
        "{ Cód. Documento: nomb-corto-proyecto_SW/HW_SRS }": "DOC-IPSEP-OSM27-SRS",
        "{ 00 }": "01",
        "{ xx/xx/xxx }": "15/06/2026",
        "{ xx/xx/xxxx }": "15/06/2026",
        "{ número de meses }": "1 mes",
        "{ Socio 1 }": "Benjamín Barbero Cortez",
        "{ Socio N }": "Pablo Gonsalez, Aldair Moyano, Carlos Rodriguez",
        "{ entidad que financia }": "la Cátedra de Aplicaciones TCP/IP",
        "{ institución }": "Facultad de Ingeniería UNRC"
    }

    replace_text_in_paragraphs(doc.paragraphs, replacements)
    replace_text_in_tables(doc.tables, replacements)

    # 2. Add Glossary terms
    # Terminos en 1.2.1, Abreviaturas en 1.2.3
    # Buscamos los parrafos
    for i, p in enumerate(doc.paragraphs):
        if "1.2.1." in p.text and "TÉRMINOS" in p.text:
            # Insert terms after this
            t1 = doc.paragraphs[i].insert_paragraph_before("Modbus TCP: Protocolo de comunicaciones de red industrial.")
            t2 = doc.paragraphs[i].insert_paragraph_before("Reconectador OSM27: Dispositivo de maniobra eléctrica fabricado por NOJA Power.")
            t3 = doc.paragraphs[i].insert_paragraph_before("ThingsBoard: Plataforma IoT para recopilación y visualización de datos.")
            t4 = doc.paragraphs[i].insert_paragraph_before("Docker: Herramienta de contenerización de software.")
            
            # We want to insert AFTER, so we have to move forward
            insert_idx = i + 1
            break
            
    # Better approach: find paragraph, then p.insert_paragraph_before is actually before. 
    # To insert after, we insert before the next paragraph.
    
    # 3. Inject HLRs
    hlrs = [
        {'codigo': 'HLR_COM_001', 'nombre': 'Soporte Modbus-TCP', 'desc': 'Relevamiento y soporte del protocolo de comunicación Modbus-TCP.', 'prioridad': 1, 'rel': 'N/A', 'trazabilidad': 'REQ-IPSEP', 'ref': 'Consigna'},
        {'codigo': 'HLR_SIM_001', 'nombre': 'Simulador OSM27', 'desc': 'Implementación de un simulador del reconectador OSM27.', 'prioridad': 1, 'rel': 'N/A', 'trazabilidad': 'REQ-IPSEP', 'ref': 'Consigna'},
        {'codigo': 'HLR_DAQ_001', 'nombre': 'Adquisición y Almacenamiento', 'desc': 'Sistema de adquisición continua y almacenamiento histórico de datos (ThingsBoard / PostgreSQL).', 'prioridad': 1, 'rel': 'HLR_COM_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'Consigna'},
        {'codigo': 'HLR_GUI_001', 'nombre': 'Dashboard Web', 'desc': 'Dashboard web de monitoreo en tiempo real e histórico.', 'prioridad': 1, 'rel': 'HLR_DAQ_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'Consigna'},
        {'codigo': 'HLR_ALR_001', 'nombre': 'Sistema de Alertas', 'desc': 'Sistema de alertas y notificaciones ante anomalías eléctricas o de conectividad.', 'prioridad': 1, 'rel': 'HLR_DAQ_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'Consigna'},
        {'codigo': 'HLR_HW_001', 'nombre': 'Requerimientos de Hardware', 'desc': 'El servidor ThingsBoard y el Gateway Modbus requerirán como mínimo: CPU de 4 Cores, 8GB de RAM y 128 GB SSD.', 'prioridad': 1, 'rel': 'N/A', 'trazabilidad': 'REQ-IPSEP', 'ref': 'N/A'},
    ]

    # Find the template table for HLR (Table 12 in list, index 11 roughly, or we just find one that has '{ Código }' in cell 1)
    template_req_table = None
    for tbl in doc.tables:
        if len(tbl.rows) == 4 and '{ Código }' in tbl.rows[0].cells[1].text:
            template_req_table = tbl
            break
            
    if not template_req_table:
        print("No se encontró la tabla de plantilla de requerimientos.")
        sys.exit(1)

    # Let's append HLR tables at the end of section 2.1
    # Find paragraph with "2.1."
    p_hlr = None
    for p in doc.paragraphs:
        if "REQUERIMIENTOS ESPECÍFICOS" in p.text:
            p_hlr = p
            break
            
    # We will insert BEFORE the "2.2. REQUERIMIENTOS ESPECIFICOS" heading
    if p_hlr:
        for hlr in hlrs:
            new_tbl = create_req_table(doc, template_req_table, hlr)
            p_hlr._p.addprevious(new_tbl)
            # Add a spacer paragraph
            from docx.oxml import OxmlElement
            spacer = OxmlElement('w:p')
            p_hlr._p.addprevious(spacer)

    # 4. Inject LLRs
    llrs = [
        {'codigo': 'LLR_DAQ_001', 'nombre': 'Lectura Modbus Básica', 'desc': 'Lectura Modbus de mediciones básicas (Tensiones, Corrientes, Frecuencia, FP).', 'prioridad': 1, 'rel': 'HLR_DAQ_001', 'trazabilidad': 'MANUAL-NOJA', 'ref': 'modbus.json'},
        {'codigo': 'LLR_DAQ_002', 'nombre': 'Lectura 32 bits', 'desc': 'Ensamblado y lectura Modbus de registros de 32 bits (Energía Acumulada).', 'prioridad': 2, 'rel': 'HLR_DAQ_001', 'trazabilidad': 'MANUAL-NOJA', 'ref': 'modbus.json'},
        {'codigo': 'LLR_DAQ_003', 'nombre': 'Lectura de Estados', 'desc': 'Polling Modbus de estados discretos (Abierto, Cerrado, Bloqueo).', 'prioridad': 1, 'rel': 'HLR_DAQ_001', 'trazabilidad': 'MANUAL-NOJA', 'ref': 'modbus.json'},
        {'codigo': 'LLR_SIM_001', 'nombre': 'Servidor Modbus', 'desc': 'Exponer servidor Modbus-TCP en puerto 502.', 'prioridad': 1, 'rel': 'HLR_SIM_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'main.py'},
        {'codigo': 'LLR_SIM_002', 'nombre': 'Variaciones Realistas', 'desc': 'Generar variaciones temporales realistas en telemetría.', 'prioridad': 2, 'rel': 'HLR_SIM_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'estado.py'},
        {'codigo': 'LLR_SIM_003', 'nombre': 'Inyección Manual', 'desc': 'Interfaz para inyectar fallas manualmente.', 'prioridad': 3, 'rel': 'HLR_SIM_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'cliente_interactivo.py'},
        {'codigo': 'LLR_GUI_001', 'nombre': 'Componentes en Tiempo Real', 'desc': 'Displays numéricos y luces de estado.', 'prioridad': 1, 'rel': 'HLR_GUI_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'Dashboard'},
        {'codigo': 'LLR_GUI_002', 'nombre': 'Gráficos Históricos', 'desc': 'Gráficos de series de tiempo históricos.', 'prioridad': 2, 'rel': 'HLR_GUI_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'Dashboard'},
        {'codigo': 'LLR_GUI_003', 'nombre': 'Panel de Eventos', 'desc': 'Panel visual de eventos y alarmas registradas.', 'prioridad': 2, 'rel': 'HLR_GUI_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'Dashboard'},
        {'codigo': 'LLR_ALR_001', 'nombre': 'Detección Backend', 'desc': 'Detección en el backend de condiciones de fallo.', 'prioridad': 1, 'rel': 'HLR_ALR_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'ThingsBoard'},
        {'codigo': 'LLR_ALR_002', 'nombre': 'Alertas Visuales', 'desc': 'Generación de alertas visuales en el Dashboard.', 'prioridad': 2, 'rel': 'HLR_ALR_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'ThingsBoard'},
        {'codigo': 'LLR_ALR_003', 'nombre': 'Notificaciones Email', 'desc': 'Envío de notificaciones por Correo Electrónico.', 'prioridad': 2, 'rel': 'HLR_ALR_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'ThingsBoard'},
        {'codigo': 'LLR_ALR_004', 'nombre': 'Notificaciones Mensajería', 'desc': 'Envío de notificaciones por Telegram o WhatsApp.', 'prioridad': 3, 'rel': 'HLR_ALR_001', 'trazabilidad': 'REQ-IPSEP', 'ref': 'ThingsBoard'},
    ]

    p_data = None
    for p in doc.paragraphs:
        if "3.  DEFINICIÓN DE DATOS" in p.text:
            p_data = p
            break

    if p_data:
        for llr in llrs:
            new_tbl = create_req_table(doc, template_req_table, llr)
            p_data._p.addprevious(new_tbl)
            spacer = OxmlElement('w:p')
            p_data._p.addprevious(spacer)

    # 5. Populate Data Definition Tables
    # Simple data: Table with 7 columns
    # Find table with header "Dato", "Abreviatura", "Descripción", "Tipo", "Un. Ing.", "Precisión", "Ref."
    table_simple = None
    for tbl in doc.tables:
        if len(tbl.columns) == 7 and "Dato" in tbl.rows[0].cells[0].text:
            table_simple = tbl
            # We want the second one (Table 27 in my dump), which is the example table. Or the first one. Let's find one that has rows.
            # Table 27 has 3 rows. We will use it.
            if len(tbl.rows) >= 2:
                table_simple = tbl
                break

    if table_simple:
        # Clear existing example rows (keep header)
        while len(table_simple.rows) > 1:
            table_simple._element.remove(table_simple.rows[1]._element)
            
        simples = [
            ("Tensión de Fase/Línea", "U", "Tensiones eléctricas", "Numérico", "kV", "3", "LLR_DAQ_001"),
            ("Corriente de Fase", "I", "Corriente eléctrica", "Numérico", "A", "1", "LLR_DAQ_001"),
            ("Potencia Activa", "P", "Potencia Activa total", "Numérico", "kW", "3", "LLR_DAQ_001"),
            ("Potencia Reactiva", "Q", "Potencia Reactiva total", "Numérico", "kVAr", "3", "LLR_DAQ_001"),
            ("Potencia Aparente", "S", "Potencia Aparente total", "Numérico", "kVA", "3", "LLR_DAQ_001"),
            ("Frecuencia", "f", "Frecuencia de red", "Numérico", "Hz", "2", "LLR_DAQ_001"),
            ("Factor de Potencia", "FP", "Factor de Potencia", "Numérico", "-", "3", "LLR_DAQ_001"),
            ("Estados/Alarmas", "EST", "Booleanos de estado", "Booleano", "-", "0", "LLR_DAQ_003"),
        ]
        for data in simples:
            new_tr = copy.deepcopy(table_simple.rows[-1]._tr)
            table_simple._tbl.append(new_tr)
            from docx.table import _Row
            new_row = _Row(new_tr, table_simple)
            for i, val in enumerate(data):
                new_row.cells[i].text = val

    # Compound data: Table with 5 columns
    table_comp = None
    for tbl in doc.tables:
        if len(tbl.columns) == 5 and "Dato Componente" in tbl.rows[0].cells[3].text:
            if len(tbl.rows) >= 2:
                table_comp = tbl
                break
                
    if table_comp:
        while len(table_comp.rows) > 1:
            table_comp._element.remove(table_comp.rows[1]._element)
            
        comps = [
            ("Energía Activa Acumulada", "E_ACT", "Acumulador 32 bits", "2 Registros Modbus 16b", "LLR_DAQ_002"),
            ("Energía Reactiva Acumulada", "E_REA", "Acumulador 32 bits", "2 Registros Modbus 16b", "LLR_DAQ_002"),
            ("Reloj Controlador", "TIME", "Timestamp de Unix", "2 Registros Modbus 16b", "LLR_DAQ_002"),
        ]
        for data in comps:
            new_tr = copy.deepcopy(table_comp.rows[-1]._tr)
            table_comp._tbl.append(new_tr)
            from docx.table import _Row
            new_row = _Row(new_tr, table_comp)
            for i, val in enumerate(data):
                new_row.cells[i].text = val

    doc.save(output_path)
    print("Documento guardado con éxito:", output_path)

if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
